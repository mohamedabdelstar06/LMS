from __future__ import annotations

from pathlib import Path
from tempfile import gettempdir
import re
import requests

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DOC = Path("C:/Users/MK/Documents/LMS/Graduation_Project_Hardware_Software_Requirements_v2.docx")
TEMP_IMG_DIR = Path(gettempdir()) / "lms_req_doc_assets"
TEMP_IMG_DIR.mkdir(parents=True, exist_ok=True)
ARCH_IMG = TEMP_IMG_DIR / "system_architecture.png"

ACCENT_BLUE = RGBColor(0, 61, 121)


def safe_name(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_-]+", "_", text).strip("_").lower()


def download_logo(name: str, url: str) -> Path | None:
    target = TEMP_IMG_DIR / f"{safe_name(name)}.png"
    try:
        response = requests.get(url, timeout=15)
        response.raise_for_status()
        target.write_bytes(response.content)
        return target
    except Exception:
        return None


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "003D79")
        tblBorders.append(elem)
    tblPr.append(tblBorders)


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT_BLUE
    return h


def add_logo_strip(doc: Document, title: str, logos: list[tuple[str, str]]):
    add_heading(doc, title, level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for name, url in logos:
        logo_path = download_logo(name, url)
        if logo_path:
            run = p.add_run()
            run.add_picture(str(logo_path), width=Inches(0.48))
            p.add_run(f" {name}   ")
        else:
            p.add_run(f"[{name}]   ")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = ACCENT_BLUE

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_borders(table)
    doc.add_paragraph("")


def set_small_margins(doc: Document):
    for section in doc.sections:
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)


def draw_box(draw: ImageDraw.ImageDraw, xy, text: str):
    draw.rectangle(xy, outline=(0, 61, 121), width=3, fill=(255, 255, 255))
    x1, y1, x2, y2 = xy
    draw.text((x1 + 12, y1 + (y2 - y1) // 2 - 8), text, fill=(0, 0, 0))


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]):
    draw.line([start, end], fill=(0, 61, 121), width=3)
    ex, ey = end
    draw.polygon([(ex, ey), (ex - 10, ey - 5), (ex - 10, ey + 5)], fill=(0, 61, 121))


def create_architecture_diagram() -> Path:
    img = Image.new("RGB", (1600, 920), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    draw_box(draw, (70, 350, 360, 500), "Flutter Web Client")
    draw_box(draw, (470, 120, 860, 260), "Auth + API Controllers")
    draw_box(draw, (470, 330, 860, 470), "Business Services Layer")
    draw_box(draw, (470, 560, 860, 700), "Background Jobs (Hangfire)")
    draw_box(draw, (970, 120, 1510, 260), "SQL Server Database")
    draw_box(draw, (970, 330, 1510, 470), "SignalR Hub")
    draw_box(draw, (970, 560, 1510, 700), "Python Data Science Pipeline")
    draw_box(draw, (970, 760, 1510, 890), "Analytics Outputs / Insights")

    draw_arrow(draw, (360, 405), (470, 190))
    draw_arrow(draw, (860, 190), (970, 190))
    draw_arrow(draw, (360, 430), (470, 395))
    draw_arrow(draw, (860, 395), (970, 395))
    draw_arrow(draw, (860, 395), (970, 625))
    draw_arrow(draw, (860, 625), (970, 625))
    draw_arrow(draw, (1240, 700), (1240, 760))

    draw.text((85, 520), "HTTPS JSON Requests / Responses", fill=(0, 61, 121))
    draw.text((930, 40), "Backend-Centric System Architecture", fill=(0, 61, 121))

    img.save(ARCH_IMG)
    return ARCH_IMG


def add_requirements_section(doc: Document):
    add_heading(doc, "Part 1: Hardware and Software Requirements", level=1)
    doc.add_paragraph(
        "This section is intentionally condensed to fit one to two pages while still covering all major technical requirements "
        "for the project: Flutter Web frontend, ASP.NET Core backend, SQL Server, and the Python Data Science layer."
    )

    add_logo_strip(
        doc,
        "Main Tools and Platforms",
        [
            ("Flutter", "https://logo.clearbit.com/flutter.dev"),
            (".NET", "https://logo.clearbit.com/dotnet.microsoft.com"),
            ("SQL Server", "https://logo.clearbit.com/microsoft.com"),
            ("Python", "https://logo.clearbit.com/python.org"),
            ("Figma", "https://logo.clearbit.com/figma.com"),
            ("scikit-learn", "https://logo.clearbit.com/scikit-learn.org"),
        ],
    )

    add_table(
        doc,
        ["Software Area", "Requirements"],
        [
            ["Frontend", "Flutter SDK 3.9.x, Dart 3.9+, flutter_bloc, dio, flutter_svg, secure storage support"],
            ["UI/UX Design", "Figma used for wireframes, user flows, design system, spacing, and component consistency"],
            ["Backend", "ASP.NET Core Web API (.NET 9), EF Core 9, JWT auth, ASP.NET Identity, SignalR, Hangfire, Serilog"],
            ["Database", "SQL Server 2019+ with migration-based schema updates and transactional consistency"],
            ["Data Science", "Python 3.10+, pandas, numpy, scikit-learn, joblib, optional shap, openpyxl/xlsxwriter"],
            ["DevOps / Tools", "Git, Swagger/OpenAPI, Postman or API testing tool, CI-friendly build pipeline"],
        ],
    )

    add_table(
        doc,
        ["Hardware Area", "Minimum", "Recommended"],
        [
            ["Development Machines", "4 cores CPU, 8 GB RAM, 256 GB SSD", "8 cores CPU, 16-32 GB RAM, 512 GB NVMe"],
            ["Backend Hosting", "4 vCPU, 8 GB RAM", "8 vCPU, 16 GB RAM + autoscaling ready"],
            ["Database Hosting", "4 vCPU, 16 GB RAM", "8 vCPU, 32 GB RAM + high-IO SSD + backup strategy"],
            ["Data Science Worker", "4 vCPU, 8 GB RAM", "8 vCPU, 16-32 GB RAM, optional GPU for heavier experiments"],
        ],
    )

    p = doc.add_paragraph()
    p.add_run("Note on Figma: ").bold = True
    p.add_run(
        "Although Figma is not part of the source code repository, it is a core project dependency for UI/UX planning, "
        "design handoff, and consistency between screens implemented in Flutter Web."
    )


def add_architecture_section(doc: Document):
    doc.add_page_break()
    add_heading(doc, "Part 2: System Architecture (Backend-Focused)", level=1)
    doc.add_paragraph(
        "This architecture section is written to fit two to three pages. The backend is the central integration point between "
        "Flutter Web clients, the SQL Server database, real-time messaging through SignalR, and analytics generated by the "
        "Data Science pipeline."
    )

    add_heading(doc, "2.1 Architecture Narrative", level=2)
    for item in [
        "Client Layer: Flutter Web sends authenticated HTTPS requests (JSON payloads) and receives structured API responses.",
        "API Layer: ASP.NET Core controllers validate incoming requests, enforce authorization policies, and route actions to services.",
        "Service Layer: business logic handles courses, users, quizzes, submissions, notifications, and analytics-trigger points.",
        "Data Layer: Entity Framework Core persists and reads normalized data from SQL Server using repositories and migration control.",
        "Realtime Layer: SignalR hub pushes live updates such as notifications and activity events to connected web clients.",
        "Async Layer: Hangfire handles scheduled and delayed tasks like digest notifications and deferred processing jobs.",
        "Data Science Integration: Python pipeline consumes exported/curated data, runs ML models, then publishes risk and insight outputs.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "2.2 Request/Response Path", level=2)
    doc.add_paragraph(
        "A typical flow starts when the Flutter Web client sends an HTTPS request with a JWT token. The backend authenticates "
        "the request, executes service-layer rules, reads/writes SQL Server through EF Core, and returns a JSON response. "
        "If the action requires live updates, the backend also emits a SignalR event. For predictive features, selected "
        "operational data is processed by the Python analytics pipeline, then summarized outputs are returned to dashboards."
    )

    add_heading(doc, "2.3 Architecture Diagram", level=2)
    diagram = create_architecture_diagram()
    doc.add_picture(str(diagram), width=Inches(10.0))
    diag_caption = doc.add_paragraph("Figure: Backend-centric architecture and request/response flow.")
    diag_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_dfd_section(doc: Document):
    doc.add_page_break()
    add_heading(doc, "Part 3: Data Flow Diagram (Text Description)", level=1)
    doc.add_paragraph(
        "This page describes the DFD in text only, as requested. It is split into approximately half a page for Level 0 and "
        "half a page for Level 1."
    )

    add_heading(doc, "3.1 DFD Level 0 (Context Level)", level=2)
    doc.add_paragraph(
        "At Level 0, SkyLearn LMS is represented as one central process that interacts with four external entities: "
        "Students, Instructors, Administrators, and the Data Science/Reporting consumer. Students submit login requests, "
        "course interactions, assignments, and quiz responses to the system. Instructors send course management data, "
        "grading actions, and content updates. Administrators provide user governance, permissions, and configuration commands. "
        "The system responds to all entities with dashboards, status updates, notifications, and academic performance results. "
        "The same central process also exchanges curated datasets and generated insights with the analytics side, enabling "
        "early-warning and performance intelligence."
    )

    add_heading(doc, "3.2 DFD Level 1 (Decomposition Level)", level=2)
    doc.add_paragraph(
        "At Level 1, the system is decomposed into major subprocesses: (1) Authentication and Access Control, "
        "(2) Learning Content and Activities Management, (3) Assessment and Evaluation, (4) Notification and Realtime "
        "Communication, and (5) Analytics and Prediction. Process (1) validates identity and role permissions before passing "
        "authorized requests forward. Process (2) handles courses, lectures, files, and participation records. Process (3) "
        "manages quizzes, assignment submissions, grading, and feedback loops. Process (4) creates and delivers notifications "
        "through API and SignalR channels. Process (5) transforms operational data into indicators such as at-risk probability, "
        "engagement segments, and performance forecasts. Data stores at this level include the primary SQL Server operational "
        "database and exported analytical datasets used by the Python ML pipeline."
    )


def build_doc() -> None:
    doc = Document()
    set_small_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)

    title = doc.add_heading("SkyLearn LMS - Graduation Book Content", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = ACCENT_BLUE

    subtitle = doc.add_paragraph("Requirements, System Architecture, and DFD Description")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = ACCENT_BLUE
    doc.add_paragraph("")

    add_requirements_section(doc)
    add_architecture_section(doc)
    add_dfd_section(doc)
    doc.save(OUTPUT_DOC)


if __name__ == "__main__":
    build_doc()
